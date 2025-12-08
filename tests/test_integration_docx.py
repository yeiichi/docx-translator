from __future__ import annotations

from pathlib import Path

import docx  # python-docx
import pytest

from docx_translator.core import translator as tr


class DummyEngine:
    """
    Simple fake TranslationEngine implementation.

    It just decorates segments so we can see that the pipeline passed them
    correctly: "<segment> -> <segment> [EN→JA]".
    """

    def __init__(self) -> None:
        self.calls: list[tuple[list[str], str, str]] = []

    def translate_segments(self, segments, source_lang: str, target_lang: str):
        seg_list = list(segments)
        self.calls.append((seg_list, source_lang, target_lang))
        return [f"{s} [{source_lang}→{target_lang}]" for s in seg_list]


def test_run_translates_segments_and_calls_writer(tmp_path: Path, monkeypatch) -> None:
    """
    High-level integration test for DocxTranslator.run with fakes.

    - Monkeypatch read_docx_to_segments to return fixed segments.
    - Monkeypatch write_segments_to_docx to capture what was written.
    - Use DummyEngine to avoid real DeepL calls.
    - Verify:
        * write_segments_to_docx receives translated segments
        * segment_mode is propagated
        * output file is created
    """
    # Arrange: fake input/output paths
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    input_path.write_bytes(b"dummy-docx-content")

    # Fake segments that the reader will "extract"
    fake_segments = ["Hello", "world"]
    captured = {"called": False, "segments": None, "mode": None}

    def fake_read_docx_to_segments(path: Path, mode: str):
        # Ensure the translator passes the right input path & mode
        assert path == input_path
        assert mode in {"run", "paragraph"}
        return list(fake_segments)

    def fake_write_segments_to_docx(
            *,
            input_path: Path,
            output_path: Path,
            translated_segments,
            mode: str,
    ):
        captured["called"] = True
        captured["segments"] = list(translated_segments)
        captured["mode"] = mode

        # For this test, just write a simple text representation so we
        # can assert that the file exists.
        output_path.write_text("\n".join(translated_segments), encoding="utf-8")

    # Patch the I/O layer inside translator module
    monkeypatch.setattr(tr, "read_docx_to_segments", fake_read_docx_to_segments)
    monkeypatch.setattr(tr, "write_segments_to_docx", fake_write_segments_to_docx)

    # Use DummyEngine instead of the real DeepL engine
    engine = DummyEngine()
    docx_translator = tr.DocxTranslator(engine=engine, segment_mode="paragraph")

    job = tr.TranslationJob(
        input_path=input_path,
        output_path=output_path,
        source_lang="EN",
        target_lang="JA",
    )

    # Act
    docx_translator.run(job)

    # Assert: I/O and engine wiring
    assert captured["called"] is True
    assert captured["mode"] == "paragraph"

    # Engine should have been called once with our fake segments
    assert len(engine.calls) == 1
    called_segments, src_lang, tgt_lang = engine.calls[0]
    assert called_segments == fake_segments
    assert src_lang == "EN"
    assert tgt_lang == "JA"

    # Writer should have seen translated segments from DummyEngine
    assert captured["segments"] == [
        "Hello [EN→JA]",
        "world [EN→JA]",
    ]
    assert output_path.is_file()
    # Optional: sanity check content
    content = output_path.read_text(encoding="utf-8").splitlines()
    assert content == captured["segments"]


def test_run_with_no_segments_copies_document(tmp_path: Path, monkeypatch) -> None:
    """
    When read_docx_to_segments returns an empty list, DocxTranslator.run should
    call write_segments_to_docx with translated_segments=[].

    The real implementation copies the document unchanged; our fake writer will
    simulate that behavior.
    """
    input_path = tmp_path / "input_empty.docx"
    output_path = tmp_path / "output_empty.docx"
    original_bytes = b"original-docx-content"
    input_path.write_bytes(original_bytes)

    def fake_read_docx_to_segments(path: Path, mode: str):
        # No segments extracted -> triggers the 'no segments' branch
        return []

    captured = {"segments": None}

    def fake_write_segments_to_docx(
            *,
            input_path: Path,
            output_path: Path,
            translated_segments,
            mode: str,
    ):
        captured["segments"] = list(translated_segments)
        # Simulate "copy document unchanged"
        output_path.write_bytes(input_path.read_bytes())

    # Patch I/O layer
    monkeypatch.setattr(tr, "read_docx_to_segments", fake_read_docx_to_segments)
    monkeypatch.setattr(tr, "write_segments_to_docx", fake_write_segments_to_docx)

    # Engine should not matter here; create a dummy that should not be called
    class NeverCalledEngine:
        def translate_segments(self, segments, source_lang, target_lang):
            raise AssertionError("Engine.translate_segments should not be called when no segments.")

    docx_translator = tr.DocxTranslator(engine=NeverCalledEngine())

    job = tr.TranslationJob(
        input_path=input_path,
        output_path=output_path,
        source_lang="EN",
        target_lang="JA",
    )

    # Act
    docx_translator.run(job)

    # Assert: writer was called with empty translated_segments and copy behavior
    assert captured["segments"] == []
    assert output_path.is_file()
    assert output_path.read_bytes() == original_bytes


class FailingEngine:
    """
    TranslationEngine implementation that always raises.

    This is used to verify that DocxTranslator.run() surfaces
    engine errors instead of silently swallowing them.
    """

    def __init__(self) -> None:
        self.calls = 0

    def translate_segments(self, segments, source_lang, target_lang):
        self.calls += 1
        raise RuntimeError("engine failure for test")


def test_run_raises_when_engine_fails(tmp_path, monkeypatch):
    """
    If the underlying translation engine raises, DocxTranslator.run()
    should surface the error (not silently continue).
    """
    # Arrange: a dummy input DOCX path
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    input_path.write_bytes(b"dummy-docx-content")

    fake_segments = ["Hello", "World"]

    # Match the real signature: read_docx_to_segments(path, mode)
    def fake_read_docx_to_segments(path: Path, mode: str):
        assert path == input_path
        assert mode in {"run", "paragraph"}
        return list(fake_segments)

    # Match the real signature: write_segments_to_docx(input_path, output_path, segments, mode)
    def fake_write_segments_to_docx(
            input_path_arg: Path,
            output_path_arg: Path,
            translated_segments,
            mode: str,
    ):
        # Should never be called because the engine fails
        pytest.fail("write_segments_to_docx should not be called when engine fails")

    monkeypatch.setattr(tr, "read_docx_to_segments", fake_read_docx_to_segments)
    monkeypatch.setattr(tr, "write_segments_to_docx", fake_write_segments_to_docx)

    failing_engine = FailingEngine()
    docx_translator = tr.DocxTranslator(engine=failing_engine)

    job = tr.TranslationJob(
        input_path=input_path,
        output_path=output_path,
        source_lang="EN",
        target_lang="JA",
    )

    # Act + Assert: engine failure should bubble up
    with pytest.raises(RuntimeError, match="engine failure for test"):
        docx_translator.run(job)

    assert failing_engine.calls == 1


@pytest.mark.slow
def test_run_real_docx_roundtrip_with_dummy_engine(tmp_path):
    """
    End-to-end smoke test using a real .docx file and DummyEngine.

    This verifies that:
    - We can read a DOCX produced by python-docx
    - Segments are passed to the engine
    - The output DOCX is written and readable
    - Basic paragraph structure is preserved

    Adjust expected strings if your segmentation is more granular
    (e.g. run-level instead of paragraph-level).
    """
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    # Create a very simple DOCX
    doc = docx.Document()
    doc.add_paragraph("Hello")
    doc.add_paragraph("World")
    doc.save(input_path)

    # Use your existing DummyEngine from this module
    docx_translator = tr.DocxTranslator(engine=DummyEngine())

    job = tr.TranslationJob(
        input_path=input_path,
        output_path=output_path,
        source_lang="EN",
        target_lang="JA",
    )

    # Act
    docx_translator.run(job)

    # Assert: output exists and contains translated text
    assert output_path.is_file()

    translated = docx.Document(output_path)
    paras = [p.text for p in translated.paragraphs]

    # If read_docx_to_segments operates at paragraph-level, this is a good expectation.
    # If it segments differently, adapt to your actual behavior.
    assert "Hello" in paras[0]
    assert "World" in paras[1]
    # And, because DummyEngine decorates segments:
    assert all("[EN→JA]" in p for p in paras if p.strip())
