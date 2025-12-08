from __future__ import annotations

from pathlib import Path
from typing import List, Union

from docx import Document

PathLike = Union[str, Path]


def write_segments_to_docx(
    input_path: PathLike,
    output_path: PathLike,
    translated_segments: List[str],
    mode: str = "run",
) -> None:
    """
    Create a translated DOCX by injecting `translated_segments` into a
    copy of the original document.

    mode:
      - "run"       : expects one segment per (paragraph.run), preserves run
                      structure and inline formatting.
      - "paragraph" : expects one segment per paragraph; sets paragraph.text
                      directly. Paragraph styles are kept, but run-level
                      formatting (bold/italic per word, etc.) may change.
    """
    in_path = Path(input_path)
    out_path = Path(output_path)

    document = Document(in_path)

    if mode == "paragraph":
        paragraphs = list(document.paragraphs)
        if len(paragraphs) != len(translated_segments):
            raise RuntimeError(
                f"Translated segments length mismatch (paragraph mode): "
                f"expected {len(paragraphs)}, got {len(translated_segments)}"
            )

        for paragraph, txt in zip(paragraphs, translated_segments):
            paragraph.text = txt or ""

        document.save(out_path)
        return

    # Default: "run" mode — existing behavior
    total_expected = sum(len(p.runs) for p in document.paragraphs)
    if total_expected != len(translated_segments):
        raise RuntimeError(
            f"Translated segments length mismatch (run mode): "
            f"expected {total_expected}, got {len(translated_segments)}"
        )

    idx = 0
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.text = translated_segments[idx]
            idx += 1

    if idx != len(translated_segments):
        raise RuntimeError(
            f"Not all translated segments were consumed: "
            f"consumed={idx}, total={len(translated_segments)}"
        )

    document.save(out_path)
