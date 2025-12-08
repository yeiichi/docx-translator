from __future__ import annotations

from pathlib import Path
from typing import List, Union

from docx import Document

PathLike = Union[str, Path]


def read_docx_to_segments(path: PathLike, mode: str = "run") -> List[str]:
    """
    Read a DOCX file and return a flat list of text segments.

    mode:
      - "run"       : one segment per run (default, best formatting preservation)
      - "paragraph" : one segment per paragraph (fewer segments, may lose
                      fine-grained inline formatting on write-back)
    """
    doc_path = Path(path)
    document = Document(doc_path)

    segments: List[str] = []

    if mode == "paragraph":
        for paragraph in document.paragraphs:
            segments.append(paragraph.text or "")
        return segments

    # Default: "run"
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            text = run.text or ""
            segments.append(text)

    return segments
