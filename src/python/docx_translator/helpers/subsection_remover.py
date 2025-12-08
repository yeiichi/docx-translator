#!/usr/bin/env python3
"""
Remove a section (heading + its subtree) from a DOCX file.
Detects headings using python-docx styles: Heading 1 .. Heading 6.

Usage:
    python subsection_remover.py INPUT.docx OUTPUT.docx
"""

import sys
from pathlib import Path
from dataclasses import dataclass
from typing import List, Optional
from docx import Document


@dataclass
class Section:
    id: int
    level: int
    title: str
    start_idx: int   # paragraph index (inclusive)
    end_idx: int     # paragraph index (exclusive)


HEADING_STYLE_NAMES = {f"Heading {i}": i for i in range(1, 7)}


def detect_sections(doc: Document) -> List[Section]:
    paras = doc.paragraphs
    raw_heads = []

    for idx, p in enumerate(paras):
        style = p.style.name if p.style else ""
        if style in HEADING_STYLE_NAMES:
            level = HEADING_STYLE_NAMES[style]
            title = p.text.strip()
            raw_heads.append((idx, level, title))

    if not raw_heads:
        return []

    sections: List[Section] = []
    n = len(raw_heads)

    for i, (start, level, title) in enumerate(raw_heads):
        # default: goes until end of document
        end = len(paras)
        # next heading with level <= this one
        for j in range(i + 1, n):
            next_start, next_level, _ = raw_heads[j]
            if next_level <= level:
                end = next_start
                break

        sections.append(Section(
            id=i + 1,
            level=level,
            title=title,
            start_idx=start,
            end_idx=end,
        ))

    return sections


def print_menu(sections: List[Section]):
    print("\nDetected DOCX sections:\n")
    for s in sections:
        indent = "  " * (s.level - 1)
        print(f"[{s.id:2}] {indent}({s.start_idx}–{s.end_idx})  Heading {s.level}: {s.title}")
    print()


def choose_section(sections: List[Section]) -> Optional[Section]:
    if not sections:
        print("No Heading 1–6 found.")
        return None

    print_menu(sections)

    while True:
        choice = input("Enter the section ID to remove (Enter = cancel): ").strip()
        if choice == "":
            return None
        if not choice.isdigit():
            print("Please enter a number.")
            continue

        id_ = int(choice)
        for s in sections:
            if s.id == id_:
                return s

        print(f"ID {id_} not found. Try again.")


def remove_section(doc: Document, section: Section):
    """Delete paragraphs from start_idx to end_idx (exclusive)."""
    paras = doc.paragraphs

    # We remove paragraphs backward to avoid index shifting
    for i in range(section.end_idx - 1, section.start_idx - 1, -1):
        p = paras[i]._element
        p.getparent().remove(p)
    return doc


def main(argv):
    if len(argv) != 3:
        print("Usage: python subsection_remover.py INPUT.docx OUTPUT.docx")
        return 1

    in_path = Path(argv[1])
    out_path = Path(argv[2])

    if not in_path.exists():
        print(f"Input not found: {in_path}")
        return 1

    doc = Document(in_path)

    sections = detect_sections(doc)
    if not sections:
        print("No DOCX headings found. Copying unchanged.")
        doc.save(out_path)
        return 0

    selected = choose_section(sections)
    if not selected:
        print("Cancelled — nothing removed.")
        return 0

    print(f"Removing section ID {selected.id} → '{selected.title}'")
    doc = remove_section(doc, selected)

    doc.save(out_path)
    print(f"Saved: {out_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
